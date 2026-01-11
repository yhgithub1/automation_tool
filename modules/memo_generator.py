# modules/memo_generator.py
import openpyxl
from docx import Document
from datetime import datetime, timedelta
import os
import sys

# Get the directory of the current script
current_dir = os.path.dirname(os.path.abspath(__file__))
# Get the project root (parent of modules directory)
project_root = os.path.dirname(current_dir)
# Add project root to Python path if not already there
if project_root not in sys.path:
    sys.path.insert(0, project_root)

from utils.file_utils import find_excel_file


def generate_memo(excel_path=None, template_path=None, output_folder=None, progress_callback=None):
    """
    生成MEMO：从Excel读取数据，为每行非空数据填充Word模板并保存
    :param excel_path: Excel文件路径（默认：tool/datasource.xlsx）
    :param template_path: Word模板路径（默认：tool/MemoTemplate.docx）
    :param output_folder: 生成文件保存文件夹路径（默认：tool/）
    :param progress_callback: 日志回调函数（传递进度到主窗口）
    :return: tuple (success: bool, message: str, generated_files: list)
    """
    # 日志发送辅助函数
    def send_log(msg):
        if progress_callback and callable(progress_callback):
            progress_callback(msg)
        print(msg)

    # 1. 初始化默认路径
    try:
        # 基础路径：桌面/tool
        tool_folder = os.path.join(os.path.expanduser("~"), "Desktop", "tool")
        if not os.path.exists(tool_folder):
            os.makedirs(tool_folder)
            send_log(f"✅ 已创建tool文件夹: {tool_folder}")

        # 默认Excel路径
        if not excel_path:
            excel_path, msg = find_excel_file()
            if not excel_path:
                send_log(f"❌ {msg}")
                return (False, msg, "")
        # 默认模板路径
        if not template_path:
            template_path = os.path.join(tool_folder, "MemoTemplate.docx")
        # 默认输出文件夹
        if not output_folder:
            output_folder = tool_folder

        send_log(f"📋 开始执行MEMO生成流程")
        send_log(f"Excel路径：{excel_path}")
        send_log(f"模板路径：{template_path}")
        send_log(f"输出文件夹：{output_folder}")

        # 2. 读取Excel数据
        send_log("\n🔍 正在读取Excel数据...")
        if not os.path.exists(excel_path):
            raise FileNotFoundError(f"Excel文件不存在：{excel_path}")

        # 读取Excel（无表头，取Sheet1工作表）
        workbook = openpyxl.load_workbook(excel_path, read_only=True)
        try:
            sheet_names = workbook.sheetnames
            send_log(f"Excel包含工作表：{sheet_names}")
            sheet = workbook['Sheet1']

            # 读取所有数据
            data = []
            for row in sheet.iter_rows(values_only=True):
                data.append(list(row))

            # Get the actual maximum column count from the sheet
            max_column = sheet.max_column

            if len(data) == 0:
                raise ValueError("Excel文件中无任何数据行")
            if max_column < 5:  # 至少需要5列（B列=1、C列=2、E列=4）
                raise ValueError(f"Excel列数不足（当前{max_column}列，需至少5列）")

            generated_files = []
            memo_count = 0

            # 循环处理每行数据
            for row_index, row in enumerate(data, start=1):
                # 检查是否为非空行（至少有序列号、公司名称、设备型号）
                sn = str(row[1]).strip() if len(row) > 1 and row[1] is not None else ""
                company_full = str(row[2]).strip() if len(row) > 2 and row[2] is not None else ""
                model = str(row[4]).strip() if len(row) > 4 and row[4] is not None else ""

                if not sn or not company_full or not model:
                    send_log(f"跳过第{row_index}行：数据不完整（序列号：{sn}，公司：{company_full}，型号：{model}）")
                    continue

                # 解析公司名称
                company_name = company_full.split('/')[-1].strip() if '/' in company_full else company_full

                send_log(f"\n📝 处理第{row_index}行数据：序列号={sn}，公司={company_name}，型号={model}")

                # 计算日期（结束日期=今天，开始日期=2天前）
                end_date = datetime.now()
                start_date = end_date - timedelta(days=2)
                excel_data = {
                    "买方": company_name,
                    "设备型号": model,
                    "序列号": sn,
                    "安装开始日期": start_date.strftime("%Y.%m.%d"),
                    "安装结束日期": end_date.strftime("%Y.%m.%d")
                }

                # 3. 填充Word模板
                doc = Document(template_path)
                keyword_mapping = {  # 关键词→数据字段的映射
                    "买方：": "买方",
                    "已完成": "设备型号",
                    "序列号：": "序列号",
                    "日期从": "安装开始日期",
                    "至": "安装结束日期"
                }
                placeholder_count = 0  # 成功替换的占位符数量

                # 处理段落中的下划线占位符
                for paragraph in doc.paragraphs:
                    for keyword, data_key in keyword_mapping.items():
                        if keyword in paragraph.text:
                            found_keyword = False
                            for run in paragraph.runs:
                                # 先找到关键词，再找后续的下划线
                                if not found_keyword and keyword in run.text:
                                    found_keyword = True
                                    continue
                                # 替换关键词后的第一个下划线
                                if found_keyword and run.underline:
                                    run.text = excel_data[data_key]
                                    placeholder_count += 1
                                    break  # 只替换第一个匹配的下划线

                # 处理表格中的下划线占位符
                for table in doc.tables:
                    for table_row in table.rows:
                        for cell in table_row.cells:
                            for paragraph in cell.paragraphs:
                                for keyword, data_key in keyword_mapping.items():
                                    if keyword in paragraph.text:
                                        found_keyword = False
                                        for run in paragraph.runs:
                                            if not found_keyword and keyword in run.text:
                                                found_keyword = True
                                                continue
                                            if found_keyword and run.underline:
                                                run.text = excel_data[data_key]
                                                placeholder_count += 1
                                                break

                # 校验替换结果
                if placeholder_count == 0:
                    raise ValueError(f"❌ 第{row_index}行未替换任何占位符！请检查模板中的关键词和下划线格式")

                # 4. 保存生成的MEMO
                output_filename = f"[{sn}]_Filled_memo.docx"
                output_path = os.path.join(output_folder, output_filename)
                doc.save(output_path)
                if not os.path.exists(output_path):
                    raise Exception(f"MEMO保存失败（文件未生成）：{output_path}")

                generated_files.append(output_path)
                memo_count += 1
                send_log(f"✅ 第{row_index}行MEMO生成成功！路径：{output_path}")

            if memo_count == 0:
                raise ValueError("未生成任何MEMO，请检查Excel数据是否完整")

            send_log(f"\n✅ 全部MEMO生成完成！共生成{memo_count}个文件")
            return (True, f"MEMO生成成功，共生成{memo_count}个文件", generated_files)
        finally:
            # 确保Excel文件被正确关闭
            workbook.close()

    except FileNotFoundError as e:
        err_msg = f"文件错误：{str(e)}"
        send_log(f"❌ {err_msg}")
        return (False, err_msg, "")
    except ValueError as e:
        err_msg = f"数据错误：{str(e)}"
        send_log(f"❌ {err_msg}")
        return (False, err_msg, "")
    except IndexError as e:
        err_msg = f"索引错误：{str(e)}（Excel数据格式可能异常）"
        send_log(f"❌ {err_msg}")
        return (False, err_msg, "")
    except Exception as e:
        err_msg = f"未知错误：{str(e)}"
        send_log(f"❌ {err_msg}")
        import traceback
        traceback.print_exc()  # 打印详细堆栈（调试用）
        return (False, err_msg, "")
# memo_generator.py 末尾添加测试代码
if __name__ == "__main__":
    # 1. 定义日志打印函数（模拟主程序的回调）
    def test_log_callback(msg):
        print(f"[测试日志] {msg}")  # 打印每一步执行日志

    # 2. 手动指定路径（避免路径问题）
    tool_folder = os.path.join(os.path.expanduser("~"), "Desktop", "tool")
    excel_path, msg = find_excel_file()
    if not excel_path:
        print(msg)
        exit(1)
    template_path = os.path.join(tool_folder, "MemoTemplate.docx")  # 模板路径
    output_folder = tool_folder  # 测试输出文件夹

    # 3. 调用MEMO生成函数（带日志回调）
    test_log_callback("开始测试MEMO生成...")
    success, message, generated_files = generate_memo(
        excel_path=excel_path,
        template_path=template_path,
        output_folder=output_folder,
        progress_callback=test_log_callback  # 传递日志函数
    )

    # 4. 打印最终结果
    test_log_callback(f"\n测试结束：")
    test_log_callback(f"是否成功：{'是' if success else '否'}")
    test_log_callback(f"结果信息：{message}")
    if success:
        test_log_callback(f"生成文件列表：")
        for file_path in generated_files:
            test_log_callback(f"  {file_path}")
