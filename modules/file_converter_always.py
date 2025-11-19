#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
备用文件转换器模块
使用纯Python库，不依赖Microsoft Office
"""

import os
import sys
import logging
import tempfile
import threading
from pathlib import Path
from datetime import datetime
from PyQt5.QtCore import QObject, pyqtSignal

try:
    import pandas as pd
    from docx import Document
    from PIL import Image, ImageDraw
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import Table, TableStyle, Paragraph, SimpleDocTemplate
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
except ImportError as e:
    logging.error(f"缺少依赖库: {e}")
    raise


class FileConverter(QObject):
    """备用文件转换器：支持Excel、Word、图片转PDF（不依赖Office）"""
    
    # 信号定义
    log_signal = pyqtSignal(str)  # 日志信号
    progress_signal = pyqtSignal(int)  # 进度信号 (0-100)
    finished_signal = pyqtSignal(bool, str)  # 完成信号 (成功/失败, 输出文件路径)
    
    def __init__(self, verbose=False):
        super().__init__()
        self.verbose = verbose
        self.logger = logging.getLogger(__name__)
        self.is_canceled = False
        
    def cancel_conversion(self):
        """取消转换任务"""
        self.is_canceled = True
        self.log_signal.emit("  正在取消转换任务...")
    
    def convert_to_pdf(self, input_file, output_file=None):
        """
        一键转换文件到PDF
        
        Args:
            input_file (str): 输入文件路径
            output_file (str, optional): 输出PDF文件路径，如果为None则自动生成
            
        Returns:
            tuple: (success: bool, output_path: str)
        """
        try:
            if self.is_canceled:
                return False, ""
                
            # 检查输入文件
            if not os.path.exists(input_file):
                error_msg = f"输入文件不存在: {input_file}"
                self.log_signal.emit(f" {error_msg}")
                return False, ""
            
            # 获取文件扩展名
            file_ext = Path(input_file).suffix.lower()
            
            # 如果没有指定输出文件，自动生成
            if not output_file:
                output_file = self._generate_output_path(input_file)
            
            # 确保输出目录存在
            os.makedirs(os.path.dirname(output_file), exist_ok=True)
            
            # 根据文件类型进行转换
            if file_ext in ['.xlsx', '.xls']:
                return self._excel_to_pdf(input_file, output_file)
            elif file_ext in ['.docx', '.doc']:
                return self._word_to_pdf(input_file, output_file)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff']:
                return self._image_to_pdf(input_file, output_file)
            else:
                error_msg = f"不支持的文件格式: {file_ext}"
                self.log_signal.emit(f" {error_msg}")
                return False, ""
                
        except Exception as e:
            error_msg = f"转换失败: {str(e)}"
            self.log_signal.emit(f" {error_msg}")
            return False, ""
    
    def _generate_output_path(self, input_file):
        """生成输出文件路径"""
        input_path = Path(input_file)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{input_path.stem}_{timestamp}.pdf"
        output_dir = os.path.join(os.path.expanduser("~"), "Desktop", "converted_pdfs")
        return os.path.join(output_dir, output_filename)
    
    def _excel_to_pdf(self, excel_file, pdf_file):
        """Excel转PDF（使用pandas + reportlab）"""
        try:
            if self.is_canceled:
                return False, ""
                
            self.log_signal.emit(f"开始转换Excel文件: {os.path.basename(excel_file)}")
            self.progress_signal.emit(10)
            
            # 读取Excel文件
            try:
                # 读取所有工作表
                excel_file_path = str(excel_file)
                if excel_file_path.endswith('.xlsx'):
                    engine = 'openpyxl'
                else:
                    engine = 'xlrd'
                
                # 读取Excel文件的所有工作表
                if excel_file_path.endswith('.xlsx'):
                    xl_file = pd.ExcelFile(excel_file_path, engine=engine)
                else:
                    xl_file = pd.ExcelFile(excel_file_path)
                
                sheets = {}
                for sheet_name in xl_file.sheet_names:
                    sheets[sheet_name] = pd.read_excel(xl_file, sheet_name=sheet_name)
                
                self.progress_signal.emit(30)
                
            except Exception as e:
                self.log_signal.emit(f" 读取Excel文件失败: {str(e)}")
                return False, ""
            
            # 创建PDF文档
            doc = SimpleDocTemplate(pdf_file, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            self.progress_signal.emit(50)
            
            # 处理每个工作表
            for sheet_name, df in sheets.items():
                # 添加工作表标题
                title = Paragraph(f"工作表: {sheet_name}", styles['Title'])
                story.append(title)
                story.append(Paragraph(f"行数: {len(df)}, 列数: {len(df.columns)}", styles['Normal']))
                story.append(Paragraph("<br/>", styles['Normal']))
                
                # 转换数据为表格格式
                data = [list(df.columns)] + df.values.tolist()
                
                # 创建表格
                table = Table(data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.black),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('PADDING', (0, 0), (-1, -1), 3),
                ]))
                
                story.append(table)
                story.append(Paragraph("<br/><br/>", styles['Normal']))
            
            # 生成PDF
            try:
                doc.build(story)
                self.progress_signal.emit(100)
                self.log_signal.emit(f" Excel转换完成: {os.path.basename(pdf_file)}")
                self.finished_signal.emit(True, pdf_file)
                return True, pdf_file
            except Exception as e:
                self.log_signal.emit(f" PDF生成失败: {str(e)}")
                return False, ""
                
        except Exception as e:
            error_msg = f"Excel转换失败: {str(e)}"
            self.log_signal.emit(f" {error_msg}")
            self.finished_signal.emit(False, "")
            return False, ""
    
    def _word_to_pdf(self, word_file, pdf_file):
        """Word转PDF（使用python-docx + reportlab）"""
        try:
            if self.is_canceled:
                return False, ""
                
            self.log_signal.emit(f" 开始转换Word文件: {os.path.basename(word_file)}")
            self.progress_signal.emit(10)
            
            # 读取Word文件
            try:
                doc = Document(word_file)
                self.progress_signal.emit(30)
                
                # 提取文本内容
                content = []
                for paragraph in doc.paragraphs:
                    content.append(paragraph.text)
                
                # 提取表格内容
                tables = []
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text)
                        table_data.append(row_data)
                    tables.append(table_data)
                
            except Exception as e:
                self.log_signal.emit(f" 读取Word文件失败: {str(e)}")
                return False, ""
            
            # 创建PDF文档
            doc = SimpleDocTemplate(pdf_file, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            self.progress_signal.emit(50)
            
            # 添加文本内容
            for text in content:
                if text.strip():
                    p = Paragraph(text, styles['Normal'])
                    story.append(p)
                    story.append(Paragraph("<br/>", styles['Normal']))
            
            # 添加表格
            for table_data in tables:
                if table_data:
                    table = Table(table_data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.black),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 10),
                        ('PADDING', (0, 0), (-1, -1), 3),
                    ]))
                    story.append(table)
                    story.append(Paragraph("<br/><br/>", styles['Normal']))
            
            # 生成PDF
            try:
                doc.build(story)
                self.progress_signal.emit(100)
                self.log_signal.emit(f" Word转换完成: {os.path.basename(pdf_file)}")
                self.finished_signal.emit(True, pdf_file)
                return True, pdf_file
            except Exception as e:
                self.log_signal.emit(f" PDF生成失败: {str(e)}")
                return False, ""
                
        except Exception as e:
            error_msg = f"Word转换失败: {str(e)}"
            self.log_signal.emit(f" {error_msg}")
            self.finished_signal.emit(False, "")
            return False, ""
    
    def _image_to_pdf(self, image_file, pdf_file):
        """图片转PDF（使用Pillow + reportlab）"""
        try:
            if self.is_canceled:
                return False, ""
                
            self.log_signal.emit(f" 开始转换图片文件: {os.path.basename(image_file)}")
            self.progress_signal.emit(20)
            
            # 使用PIL处理图片
            img = Image.open(image_file)
            
            # 获取图片尺寸
            img_width, img_height = img.size
            
            # 使用A4纸张大小，图片自适应
            pdf_width, pdf_height = A4  # 固定使用A4纸张大小
            
            # 如果图片是横向的，使用横向A4
            if img_width > img_height:
                pdf_width, pdf_height = A4[1], A4[0]  # 横向A4
            
            # 计算图片在PDF中的显示尺寸
            margin = 50
            max_width = pdf_width - 2 * margin
            max_height = pdf_height - 2 * margin
            
            # 保持宽高比缩放
            scale = min(max_width / img_width, max_height / img_height, 1.0)
            display_width = img_width * scale
            display_height = img_height * scale
            
            # 创建PDF
            c = canvas.Canvas(pdf_file, pagesize=(pdf_width, pdf_height))
            self.progress_signal.emit(60)
            
            # 在页面居中绘制图片
            x = (pdf_width - display_width) / 2
            y = (pdf_height - display_height) / 2
            
            # 如果是PNG且有透明背景，需要处理
            if img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info):
                # 创建临时RGB图片
                if img.mode == 'P':
                    img = img.convert('RGBA')
                img_rgb = Image.new('RGB', img.size, (255, 255, 255))
                img_rgb.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else img.info.get('transparency'))
                img = img_rgb
            
            # 保存临时图片文件
            temp_path = None
            try:
                with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                    temp_path = temp_file.name
                
                # 转换为RGB模式确保兼容性
                if img.mode not in ('RGB', 'RGBA'):
                    img = img.convert('RGB')
                img.save(temp_path, 'JPEG', quality=95)
                
                # 确保文件已保存
                import time
                time.sleep(0.1)  # 短暂延迟确保文件写入完成
                
                c.drawImage(temp_path, x, y, display_width, display_height)
                
            finally:
                # 关闭图片文件
                img.close()
                # 清理临时文件
                if temp_path and os.path.exists(temp_path):
                    try:
                        os.unlink(temp_path)
                    except:
                        pass  # 忽略删除错误
            
            c.showPage()
            c.save()
            
            self.progress_signal.emit(100)
            self.log_signal.emit(f" 图片转换完成: {os.path.basename(pdf_file)}")
            self.finished_signal.emit(True, pdf_file)
            return True, pdf_file
            
        except Exception as e:
            error_msg = f"图片转换失败: {str(e)}"
            self.log_signal.emit(f" {error_msg}")
            self.finished_signal.emit(False, "")
            return False, ""
    
    def batch_convert(self, input_files, output_dir=None):
        """
        批量转换文件
        
        Args:
            input_files (list): 输入文件路径列表
            output_dir (str, optional): 输出目录，如果为None则使用桌面converted_pdfs文件夹
            
        Returns:
            tuple: (success_count: int, failed_count: int, results: list)
        """
        try:
            if not input_files:
                self.log_signal.emit(" 没有文件需要转换")
                return 0, 0, []
            
            if self.is_canceled:
                return 0, 0, []
            
            # 设置输出目录
            if not output_dir:
                output_dir = os.path.join(os.path.expanduser("~"), "Desktop", "converted_pdfs")
            os.makedirs(output_dir, exist_ok=True)
            
            self.log_signal.emit(f" 输出目录: {output_dir}")
            self.log_signal.emit(f" 开始批量转换 {len(input_files)} 个文件...")
            
            success_count = 0
            failed_count = 0
            results = []
            
            for i, input_file in enumerate(input_files):
                if self.is_canceled:
                    self.log_signal.emit("  批量转换已取消")
                    break
                
                # 计算进度
                progress = int((i / len(input_files)) * 100)
                self.progress_signal.emit(progress)
                
                # 生成输出文件路径
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                input_path = Path(input_file)
                output_file = os.path.join(output_dir, f"{input_path.stem}_{timestamp}.pdf")
                
                # 执行转换
                success, output_path = self.convert_to_pdf(input_file, output_file)
                
                if success:
                    success_count += 1
                    results.append((input_file, output_path, True))
                    self.log_signal.emit(f" {i+1}/{len(input_files)} 转换成功: {os.path.basename(output_path)}")
                else:
                    failed_count += 1
                    results.append((input_file, output_path, False))
                    self.log_signal.emit(f" {i+1}/{len(input_files)} 转换失败: {os.path.basename(input_file)}")
            
            # 完成
            self.progress_signal.emit(100)
            self.log_signal.emit(f" 批量转换完成！成功: {success_count}, 失败: {failed_count}")
            self.finished_signal.emit(success_count > 0, output_dir)
            
            return success_count, failed_count, results
            
        except Exception as e:
            error_msg = f"批量转换失败: {str(e)}"
            self.log_signal.emit(f" {error_msg}")
            return 0, len(input_files), []


def main():
    """独立测试函数"""
    import sys
    
    if len(sys.argv) < 2:
        print("使用方法: python file_converter_always.py <输入文件路径> [输出文件路径]")
        print("支持的文件格式: Excel(.xlsx/.xls), Word(.docx/.doc), 图片(.jpg/.png/.bmp/.gif/.tiff)")
        return
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    # 创建转换器
    converter = FileConverter()
    
    # 设置日志回调
    def log_callback(msg):
        print(f"[日志] {msg}")
    
    converter.log_signal.connect(log_callback)
    
    # 执行转换
    print(f"开始转换文件: {input_file}")
    success, output_path = converter.convert_to_pdf(input_file, output_file)
    
    if success:
        print(f"转换成功！输出文件: {output_path}")
    else:
        print("转换失败！")


if __name__ == "__main__":
    main()
